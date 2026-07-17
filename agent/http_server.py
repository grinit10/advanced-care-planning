"""
HTTP API server for session management.

Runs alongside the LiveKit agent on a separate port (8082).
Provides endpoints for the frontend to manage ACP sessions.

Endpoints:
  GET   /health                        — Health check
  GET   /events/{room_id}              — SSE stream for live preference/plan updates
  GET   /plan/{room_id}                — Get session data (transcript, preferences, summary)
  POST  /email/{room_id}               — Register an email address  {email: "..."}
  POST  /send-plan/{room_id}           — Email the plan to all registered addresses
  POST  /close/{room_id}               — Close session, delete data + audio
"""

import asyncio
import hashlib
import io
import json
import logging
import os

import docx
from aiohttp import web
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from email_sender import is_configured as email_configured
from email_sender import send_plan_email
from session_store import SessionStore

# CORS headers — allow the frontend dev server to call this API
_CORS_HEADERS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}


@web.middleware
async def cors_middleware(request: web.Request, handler):
    """Add CORS headers to every response, and handle OPTIONS preflight."""
    if request.method == "OPTIONS":
        return web.Response(headers=_CORS_HEADERS)
    resp = await handler(request)
    for key, val in _CORS_HEADERS.items():
        resp.headers[key] = val
    return resp


logger = logging.getLogger("acp-agent.http")

# Global reference set by the agent at startup
_store: SessionStore | None = None


def init(store: SessionStore):
    """Initialise the HTTP server with a reference to the session store."""
    global _store
    _store = store


async def _get_store() -> SessionStore:
    assert _store is not None, "HTTP server not initialised with session store"
    return _store


# --- Helper: generate plan summary from pre-extracted preferences ---


def _pref_value(val) -> str:
    """Format a preference value for display."""
    if val is None:
        return "Not yet discussed"
    if isinstance(val, bool):
        return "Yes" if val else "No"
    if isinstance(val, list):
        if not val:
            return "None specified"
        if len(val) == 1:
            return str(val[0])
        return ", ".join(str(v) for v in val)
    if isinstance(val, str):
        return val.strip() or "Not yet discussed"
    return str(val)


def _build_summary_from_prefs(prefs: dict) -> str:
    """Build a rich plain-text summary from pre-extracted preferences."""
    parts = [
        "Here is a summary of what we discussed during your Advanced Care Planning conversation.\n"
    ]

    # Substitute Decision-Maker
    sdm = prefs.get("substitute_decision_maker", {})
    if sdm.get("discussed"):
        parts.append("SUBSTITUTE DECISION-MAKER")
        if sdm.get("name"):
            parts.append(f"  Who: {sdm['name']}")
        if sdm.get("relationship"):
            parts.append(f"  Relationship: {sdm['relationship']}")
        if sdm.get("notes"):
            parts.append(f"  Notes: {sdm['notes']}")
        parts.append("")

    # Quality of Life
    qol = prefs.get("quality_of_life", {})
    if qol.get("discussed"):
        parts.append("QUALITY OF LIFE")
        vals = qol.get("values", [])
        if vals:
            parts.append(f"  What matters: {_pref_value(vals)}")
        fears = qol.get("fears", [])
        if fears:
            parts.append(f"  Concerns: {_pref_value(fears)}")
        if qol.get("notes"):
            parts.append(f"  Notes: {qol['notes']}")
        parts.append("")

    # Treatment Preferences
    tx = prefs.get("treatment_preferences", {})
    if tx.get("discussed"):
        parts.append("TREATMENT PREFERENCES")
        for key, label in [
            ("life_support", "Life support"),
            ("cpr", "CPR"),
            ("feeding_tubes", "Feeding tubes"),
            ("pain_management", "Pain management"),
        ]:
            val = tx.get(key)
            if val is not None:
                parts.append(f"  {label}: {_pref_value(val)}")
        if tx.get("notes"):
            parts.append(f"  Notes: {tx['notes']}")
        parts.append("")

    # Personal Beliefs
    beliefs = prefs.get("personal_beliefs", {})
    if beliefs.get("discussed"):
        parts.append("PERSONAL VALUES AND BELIEFS")
        if beliefs.get("faith_role"):
            parts.append(f"  Faith/Spirituality: {beliefs['faith_role']}")
        cult = beliefs.get("cultural_values", [])
        if cult:
            parts.append(f"  Cultural values: {_pref_value(cult)}")
        if beliefs.get("notes"):
            parts.append(f"  Notes: {beliefs['notes']}")
        parts.append("")

    # Specific Scenarios
    scenarios = prefs.get("specific_scenarios", {})
    if scenarios.get("discussed"):
        parts.append("SPECIFIC SCENARIOS")
        for key, label in [
            ("dementia", "Advanced dementia"),
            ("coma", "Persistent coma"),
            ("terminal_illness", "Terminal illness"),
        ]:
            val = scenarios.get(key)
            if val is not None:
                parts.append(f"  {label}: {_pref_value(val)}")
        if scenarios.get("notes"):
            parts.append(f"  Notes: {scenarios['notes']}")
        parts.append("")

    # Dignity & Values
    dignity = prefs.get("dignity_and_values", {})
    if dignity.get("discussed"):
        parts.append("DIGNITY AND VALUES")
        meaning = dignity.get("meaning_of_life", [])
        if meaning:
            parts.append(f"  What gives life meaning: {_pref_value(meaning)}")
        if dignity.get("dignity_definition"):
            parts.append(f"  Definition of dignity: {dignity['dignity_definition']}")
        if dignity.get("notes"):
            parts.append(f"  Notes: {dignity['notes']}")
        parts.append("")

    # Check if anything was discussed
    if len(parts) == 1:
        parts.append(
            "The conversation was still in its early stages. "
            "No specific preferences were captured yet."
        )

    return "\n".join(parts)


async def _generate_summary(room_id: str) -> str:
    """Generate a plan summary from pre-extracted preferences."""
    store = await _get_store()
    transcript = await store.get_transcript(room_id)

    if not transcript:
        return "The conversation was too short to generate a summary."

    # Use pre-extracted preferences if available, fall back to transcript
    prefs = await store.get_preferences_json(room_id)
    if prefs:
        summary = _build_summary_from_prefs(prefs)
    else:
        summary = (
            "The conversation was exploratory in nature. "
            "Topics included general values and preferences for future healthcare."
        )

    summary += (
        f"\n\nThe conversation had {len(transcript)} exchanges. "
        "Your full transcript is included below."
    )

    return summary


# --- Route handlers ---


async def handle_health(request: web.Request) -> web.Response:
    return web.json_response({"status": "ok"})


async def handle_events(request: web.Request) -> web.Response:
    """Server-Sent Events endpoint for live preference/plan updates.

    Instead of the frontend polling every 1s, this endpoint streams
    events when data actually changes. The polling still happens but
    server-side (every 3s), and only sends data when the content hash
    differs from the last push. Uses a single HTTP connection that
    the browser's EventSource API auto-reconnects if dropped.
    """
    room_id = request.match_info["room_id"]
    store = await _get_store()

    response = web.StreamResponse(
        status=200,
        reason="OK",
        headers={
            "Content-Type": "text/event-stream",
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Access-Control-Allow-Origin": "*",
            "X-Accel-Buffering": "no",
        },
    )
    await response.prepare(request)

    last_prefs_hash = ""
    last_summary_hash = ""
    poll_interval = 3  # seconds

    try:
        while True:
            # Check preferences
            prefs = await store.get_preferences_json(room_id)
            prefs_hash = (
                hashlib.sha256(
                    json.dumps(prefs, sort_keys=True, default=str).encode()
                ).hexdigest()
                if prefs
                else ""
            )

            if prefs_hash and prefs_hash != last_prefs_hash:
                last_prefs_hash = prefs_hash
                data = json.dumps(
                    {
                        "type": "preferences",
                        "preferences": prefs,
                    }
                )
                await response.write(f"event: preferences\ndata: {data}\n\n".encode())

            # Check plan summary
            summary = await store.get_plan_summary(room_id)
            summary_hash = hashlib.sha256((summary or "").encode()).hexdigest()

            if summary and summary_hash != last_summary_hash:
                last_summary_hash = summary_hash
                data = json.dumps(
                    {
                        "type": "plan_summary",
                        "summary": summary,
                    }
                )
                await response.write(f"event: plan_summary\ndata: {data}\n\n".encode())

            # Send keepalive comment every poll cycle
            await response.write(b": keepalive\n\n")
            await response.drain()

            await asyncio.sleep(poll_interval)
    except (asyncio.CancelledError, ConnectionResetError, ConnectionAbortedError):
        logger.debug("SSE connection closed for room %s", room_id)
    return response


async def handle_get_recording(request: web.Request) -> web.Response:
    """Serve the conversation audio recording as a downloadable WAV file."""
    room_id = request.match_info["room_id"]
    store = await _get_store()
    audio_path = await store.get_audio_path(room_id)

    if not audio_path or not os.path.exists(audio_path):
        return web.json_response({"error": "Recording not available"}, status=404)

    filename = f"acp-conversation-{room_id}.wav"
    return web.FileResponse(
        path=audio_path,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "audio/wav",
        },
    )


async def handle_get_transcript_download(request: web.Request) -> web.Response:
    """Return the transcript as plain text for download."""
    room_id = request.match_info["room_id"]
    store = await _get_store()
    transcript = await store.get_transcript(room_id)

    if not transcript:
        return web.json_response({"error": "Transcript not available"}, status=404)

    lines = []
    for entry in transcript:
        role = "You" if entry.get("role") == "user" else "Assistant"
        text = entry.get("text", "")
        lines.append(f"[{role}] {text}")

    text = "\n\n".join(lines)
    filename = f"acp-transcript-{room_id}.txt"
    return web.Response(
        body=text,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "text/plain; charset=utf-8",
        },
    )


async def handle_get_plan_docx(request: web.Request) -> web.Response:
    """Generate and return the ACP plan as a formatted Word Document (.docx)."""
    room_id = request.match_info["room_id"]
    store = await _get_store()

    transcript = await store.get_transcript(room_id)
    preferences = await store.get_preferences_json(room_id)
    summary = await store.get_plan_summary(room_id)

    if not summary:
        summary = await _generate_summary(room_id)
        await store.set_plan_summary(room_id, summary)

    doc = docx.Document()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Advanced Care Planning Summary")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(44, 95, 124)  # #2c5f7c
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    h = doc.add_paragraph()
    hrun = h.add_run("Plan Overview")
    hrun.font.name = "Arial"
    hrun.font.size = Pt(14)
    hrun.font.bold = True
    hrun.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph(summary)

    # Preferences
    h = doc.add_paragraph()
    hrun = h.add_run("Your Preferences")
    hrun.font.name = "Arial"
    hrun.font.size = Pt(14)
    hrun.font.bold = True
    hrun.font.color.rgb = RGBColor(51, 51, 51)

    sections = [
        (
            "substitute_decision_maker",
            "Substitute Decision-Maker",
            [("name", "Who"), ("relationship", "Relationship")],
        ),
        (
            "quality_of_life",
            "Quality of Life",
            [("values", "What matters"), ("fears", "Concerns")],
        ),
        (
            "treatment_preferences",
            "Treatment Preferences",
            [
                ("life_support", "Life support"),
                ("cpr", "CPR"),
                ("feeding_tubes", "Feeding tubes"),
                ("pain_management", "Pain management"),
            ],
        ),
        (
            "personal_beliefs",
            "Values & Beliefs",
            [
                ("faith_role", "Faith/Spirituality"),
                ("cultural_values", "Cultural values"),
            ],
        ),
        (
            "specific_scenarios",
            "Specific Scenarios",
            [
                ("dementia", "Dementia"),
                ("coma", "Coma"),
                ("terminal_illness", "Terminal illness"),
            ],
        ),
        (
            "dignity_and_values",
            "Dignity & Values",
            [
                ("meaning_of_life", "What gives life meaning"),
                ("dignity_definition", "Definition of dignity"),
            ],
        ),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Preference Category"
    hdr_cells[1].text = "Recorded Preference"
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = "Arial"

    for sec_key, sec_label, fields in sections:
        sec_data = preferences.get(sec_key, {})
        if not sec_data or not sec_data.get("discussed"):
            row_cells = table.add_row().cells
            row_cells[0].text = sec_label
            row_cells[1].text = "Not yet discussed"
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
            continue

        row_cells = table.add_row().cells
        row_cells[0].text = sec_label
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = ""
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"

        for f_key, f_label in fields:
            raw_val = sec_data.get(f_key)
            if raw_val is not None and raw_val != "":
                val_str = _pref_value(raw_val)
                if val_str.strip():
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"  • {f_label}"
                    row_cells[1].text = val_str
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Arial"

        notes = sec_data.get("notes")
        if notes:
            val_str = _pref_value(notes)
            if val_str.strip():
                row_cells = table.add_row().cells
                row_cells[0].text = "  • Notes"
                row_cells[1].text = val_str
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Arial"

    # Transcript
    h = doc.add_paragraph()
    doc.add_paragraph()
    hrun = h.add_run("Conversation Transcript")
    hrun.font.name = "Arial"
    hrun.font.size = Pt(14)
    hrun.font.bold = True
    hrun.font.color.rgb = RGBColor(51, 51, 51)

    if transcript:
        for entry in transcript:
            p = doc.add_paragraph()
            role_label = "You" if entry.get("role") == "user" else "Assistant"
            text = entry.get("text", "")
            role_run = p.add_run(f"{role_label}: ")
            role_run.font.name = "Arial"
            role_run.bold = True
            text_run = p.add_run(text)
            text_run.font.name = "Arial"
    else:
        p = doc.add_paragraph("No transcript available.")
        p.runs[0].font.name = "Arial"

    # Footer Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p_run = p.add_run(
        "Next Steps: Share these wishes with your substitute decision-maker, family, and GP. Formalise your Advance Care Directive through your state or territory's health department."
    )
    p_run.font.name = "Arial"
    p_run.font.size = Pt(10)
    p_run.font.italic = True

    p2 = doc.add_paragraph()
    p2_run = p2.add_run(
        "This summary was generated by an AI assistant and should be reviewed by you and your healthcare provider before being formalised. It is not a legal document."
    )
    p2_run.font.name = "Arial"
    p2_run.font.size = Pt(9)
    p2_run.font.italic = True
    p2_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save to BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"acp-plan-{room_id}.docx"
    return web.Response(
        body=bio.read(),
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
    )


async def handle_get_transcript_json(request: web.Request) -> web.Response:
    """Return the transcript as JSON for live polling."""
    room_id = request.match_info["room_id"]
    store = await _get_store()
    transcript = await store.get_transcript(room_id)
    return web.json_response(
        {
            "room_id": room_id,
            "transcript": transcript or [],
        }
    )


async def handle_get_preferences(request: web.Request) -> web.Response:
    """Return the live-extracted preferences JSON for a session."""
    room_id = request.match_info["room_id"]
    store = await _get_store()
    prefs = await store.get_preferences_json(room_id)
    return web.json_response(
        {
            "room_id": room_id,
            "preferences": prefs if prefs else {},
        }
    )


async def handle_get_plan(request: web.Request) -> web.Response:
    room_id = request.match_info["room_id"]
    store = await _get_store()
    status = await store.get_status(room_id)

    if not status:
        return web.json_response({"error": "Session not found"}, status=404)

    transcript = await store.get_transcript(room_id)
    preferences = await store.get_preferences(room_id)
    summary = await store.get_plan_summary(room_id)

    if not summary:
        summary = await _generate_summary(room_id)
        await store.set_plan_summary(room_id, summary)

    return web.json_response(
        {
            "room_id": room_id,
            "status": status,
            "summary": summary,
            "preferences": preferences,
            "transcript": transcript[-50:],  # last 50 entries
            "email_count": len(await store.get_emails(room_id)),
        }
    )


async def handle_add_email(request: web.Request) -> web.Response:
    room_id = request.match_info["room_id"]
    store = await _get_store()

    try:
        body = await request.json()
    except Exception:
        return web.json_response({"error": "Invalid JSON"}, status=400)

    email = body.get("email", "").strip().lower()
    if not email or "@" not in email:
        return web.json_response({"error": "Valid email address required"}, status=400)

    await store.add_email(room_id, email)
    email_count = len(await store.get_emails(room_id))
    logger.info("Email added for session %s: %s", room_id, email)

    return web.json_response(
        {
            "status": "ok",
            "email": email,
            "email_count": email_count,
            "message": "Email address registered. Click 'Send Plan' when you're ready.",
        }
    )


async def handle_send_plan(request: web.Request) -> web.Response:
    room_id = request.match_info["room_id"]
    store = await _get_store()

    if not email_configured():
        return web.json_response(
            {
                "error": (
                    "Email service not configured. "
                    "Set ACS_CONNECTION_STRING and ACS_SENDER_DOMAIN in .env"
                ),
            },
            status=500,
        )

    emails = await store.get_emails(room_id)
    if not emails:
        return web.json_response(
            {
                "error": "No email addresses registered. Add an email first.",
            },
            status=400,
        )

    # Get session data
    transcript = await store.get_transcript(room_id)
    preferences = await store.get_preferences(room_id)
    summary = await store.get_plan_summary(room_id)
    audio_path = await store.get_audio_path(room_id)

    if not summary:
        summary = await _generate_summary(room_id)
        await store.set_plan_summary(room_id, summary)

    # Send to all registered emails
    results = []
    for email in emails:
        success = send_plan_email(
            to_email=email,
            plan_summary=summary,
            transcript=transcript,
            preferences=preferences,
            audio_path=audio_path,
        )
        results.append({"email": email, "success": success})

    success_count = sum(1 for r in results if r["success"])
    logger.info(
        "Plan sent for session %s: %d/%d emails delivered",
        room_id,
        success_count,
        len(results),
    )

    return web.json_response(
        {
            "status": "partial" if success_count < len(results) else "sent",
            "results": results,
            "message": (f"Plan sent to {success_count} of {len(results)} email(s)."),
        }
    )


async def handle_close(request: web.Request) -> web.Response:
    room_id = request.match_info["room_id"]
    store = await _get_store()

    status = await store.get_status(room_id)
    if not status:
        return web.json_response({"error": "Session not found"}, status=404)

    if status == "closed":
        return web.json_response({"error": "Session already closed"}, status=400)

    # Get audio path for cleanup
    audio_path = await store.get_audio_path(room_id)

    # Send plan to registered emails before deleting data
    emails = await store.get_emails(room_id)
    if emails and email_configured():
        transcript = await store.get_transcript(room_id)
        preferences = await store.get_preferences(room_id)
        summary = await store.get_plan_summary(room_id)
        if not summary:
            summary = await _generate_summary(room_id)
        for email in emails:
            try:
                send_plan_email(
                    to_email=email,
                    plan_summary=summary,
                    transcript=transcript,
                    preferences=preferences,
                    audio_path=audio_path,
                )
            except Exception as e:
                logger.error("Failed to send email to %s on close: %s", email, e)

    # Close session and delete Redis data
    data = await store.close_session(room_id)

    # Delete audio file
    if audio_path:
        from audio_recorder import ConversationRecorder

        ConversationRecorder.cleanup(audio_path)

    transcript_count = len(data.get("transcript", []))
    logger.info(
        "Session closed: %s (%d transcript entries, audio deleted: %s)",
        room_id,
        transcript_count,
        "yes" if audio_path else "no",
    )

    return web.json_response(
        {
            "status": "closed",
            "room_id": room_id,
            "message": "Session closed. Your data has been deleted.",
        }
    )


# --- App factory ---


def create_app() -> web.Application:
    """Create the aiohttp application."""
    app = web.Application(middlewares=[cors_middleware])
    app.router.add_get("/health", handle_health)
    app.router.add_get("/events/{room_id}", handle_events)
    app.router.add_get("/preferences/{room_id}", handle_get_preferences)
    app.router.add_get("/plan/{room_id}", handle_get_plan)
    app.router.add_get("/recording/{room_id}", handle_get_recording)
    app.router.add_get("/transcript/{room_id}", handle_get_transcript_download)
    app.router.add_get("/plan-docx/{room_id}", handle_get_plan_docx)
    app.router.add_get("/transcript-json/{room_id}", handle_get_transcript_json)
    app.router.add_post("/email/{room_id}", handle_add_email)
    app.router.add_post("/send-plan/{room_id}", handle_send_plan)
    app.router.add_post("/close/{room_id}", handle_close)
    return app


async def run_server(host: str = "0.0.0.0", port: int = 8082):
    """Start the HTTP server."""
    app = create_app()
    runner = web.AppRunner(app)
    await runner.setup()
    site = web.TCPSite(runner, host, port)
    await site.start()
    logger.info("HTTP API server running on http://%s:%d", host, port)
    # Keep running
    await asyncio.Event().wait()
